import os
import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_DIR = "./data"
OUT_DIR  = "./docx"


# ── Helpers ────────────────────────────────────────────────────────────────

def load_yaml(filename):
    with open(os.path.join(DATA_DIR, filename), encoding="utf-8") as f:
        return yaml.safe_load(f)


def t(obj, lang):
    if isinstance(obj, dict):
        return obj.get(lang) or obj.get("en") or ""
    return obj or ""


def sfont(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text.upper())
    sfont(r, size=11, bold=True, color=(31, 73, 125))
    add_hr(doc)


def two_col_row(doc, left, right="", left_bold=True, left_size=11, right_size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(left)
    sfont(r1, bold=left_bold, size=left_size)
    if right:
        p.add_run("\t")
        r2 = p.add_run(right)
        sfont(r2, italic=True, size=right_size, color=(89, 89, 89))
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr.append(tabs)


def subtitle_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    sfont(r, italic=True, size=10, color=(89, 89, 89))


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent        = Inches(0.25)
    p.paragraph_format.first_line_indent  = Inches(-0.18)
    p.paragraph_format.space_after        = Pt(2)
    r = p.add_run(text)
    sfont(r, size=10.5)


def plain(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    sfont(r, size=size)


def kv_line(doc, key, value):
    p  = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rk = p.add_run(f"{key}: ")
    sfont(rk, bold=True, size=10.5)
    rv = p.add_run(value)
    sfont(rv, size=10.5)


# ── Section renderers ──────────────────────────────────────────────────────

def render_header(doc, data, lang):
    name = t(data.get("name", ""), lang)
    p    = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    sfont(p.add_run(name), size=20, bold=True)

    contact = data.get("contact", {})
    parts   = []
    for field in ("phone", "email", "location", "linkedin", "github", "website"):
        item = contact.get(field)
        if not item:
            continue
        val = item.get("value", "") if isinstance(item, dict) else str(item)
        if val:
            parts.append(val)

    if parts:
        p2 = doc.add_paragraph(" | ".join(parts))
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        for r in p2.runs:
            sfont(r, size=10)


def render_summary(doc, data, lang, labels):
    raw = data.get("professional_summary")
    if not raw:
        return
    text = t(raw, lang).strip()
    if not text:
        return
    section_heading(doc, labels.get("summary", "Summary"))
    plain(doc, text)


def render_skills(doc, skills_data, lang, labels):
    skills = skills_data.get("skills")
    if not skills:
        return
    section_heading(doc, labels.get("skills", "Skills"))
    for key, cat in skills.items():
        cat_name  = t(cat.get("name", key), lang)
        subskills = [t(s, lang) for s in cat.get("subskills", []) if t(s, lang)]
        if subskills:
            kv_line(doc, cat_name, ", ".join(subskills))


def render_experience(doc, exp_data, lang, labels):
    experience = exp_data.get("experience")
    if not experience:
        return
    section_heading(doc, labels.get("experience", "Experience"))

    for job in experience:
        company  = t(job.get("company",    ""), lang)
        location = t(job.get("location",   ""), lang)
        start    = t(job.get("start_date", ""), lang)
        end      = t(job.get("end_date",   ""), lang)
        dates    = f"{start} - {end}" if start or end else ""
        right    = f"{location}  |  {dates}" if location and dates else location or dates

        two_col_row(doc, company, right, left_bold=True, left_size=11)

        positions = job.get("positions", [])
        for pos in positions:
            title     = t(pos.get("job_title",   ""), lang)
            pos_start = t(pos.get("start_date",  ""), lang)
            pos_end   = t(pos.get("end_date",    ""), lang)
            pos_dates = f"{pos_start} - {pos_end}" if pos_start or pos_end else ""

            if len(positions) > 1:
                two_col_row(doc, title, pos_dates, left_bold=False, left_size=10.5)
            else:
                subtitle_line(doc, title)

            for resp in pos.get("responsibilities", []):
                bullet(doc, t(resp, lang))


def render_education(doc, edu_data, lang, labels):
    items = edu_data.get("education", [])
    if not items:
        return
    section_heading(doc, labels.get("education", "Education"))
    for edu in items:
        degree  = t(edu.get("degree", ""), lang)
        year    = str(edu.get("year", ""))
        inst    = edu.get("institution", "")
        loc     = edu.get("location", "")
        two_col_row(doc, degree, year)
        parts = [p for p in [inst, loc] if p]
        if parts:
            subtitle_line(doc, "  |  ".join(parts))


def render_certifications(doc, cert_data, lang, labels):
    items = cert_data.get("certifications", [])
    if not items:
        return
    section_heading(doc, labels.get("certifications", "Awards & Certifications"))
    for item in items:
        name   = t(item.get("name", ""), lang)
        date   = t(item.get("date", ""), lang)
        issuer = item.get("issuer", "")
        two_col_row(doc, name, date)
        if issuer:
            subtitle_line(doc, issuer)


def render_projects(doc, proj_data, lang, labels):
    items = proj_data.get("projects", [])
    if not items:
        return
    section_heading(doc, labels.get("projects", "Projects"))
    for proj in items:
        name  = proj.get("name", "")
        url   = proj.get("url",  "")
        desc  = t(proj.get("description", ""), lang)
        techs = proj.get("technologies", [])
        two_col_row(doc, name, url)
        if desc:
            plain(doc, desc)
        if techs:
            kv_line(doc, "Technologies", ", ".join(techs))


def render_languages(doc, lang_data, lang, labels):
    items = lang_data.get("languages", {})
    if not items:
        return
    section_heading(doc, labels.get("languages", "Languages"))
    if isinstance(items, dict):
        for key, val in items.items():
            name  = key.capitalize()
            prof  = t(val.get("proficiency", ""), lang)
            kv_line(doc, name, prof)
    else:
        for item in items:
            name = t(item.get("language", item.get("name", "")), lang)
            prof = t(item.get("proficiency", ""), lang)
            kv_line(doc, name, prof)


# ── Builder ────────────────────────────────────────────────────────────────

def build_resume(lang, main_data, exp_data, edu_data, skills_data,
                 lang_data, cert_data, proj_data, settings):

    labels_raw = settings.get("section_labels", {})
    labels     = {k: t(v, lang) for k, v in labels_raw.items()}

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.85)
        sec.right_margin  = Inches(0.85)

    render_header(doc, main_data, lang)
    render_summary(doc, main_data, lang, labels)
    render_skills(doc, skills_data, lang, labels)
    render_experience(doc, exp_data, lang, labels)
    render_education(doc, edu_data, lang, labels)
    render_certifications(doc, cert_data, lang, labels)
    render_projects(doc, proj_data, lang, labels)
    render_languages(doc, lang_data, lang, labels)

    name_str = t(main_data.get("name", {}), lang).replace(" ", "_")
    out_file  = os.path.join(OUT_DIR, f"resume_{lang}_{name_str}.docx")
    doc.save(out_file)
    print(f"Saved: {out_file}")
    return out_file


# ── Main ───────────────────────────────────────────────────────────────────

def main():
    settings    = load_yaml("settings.yaml")
    main_data   = load_yaml("data.yaml")
    exp_data    = load_yaml("experience.yaml")
    edu_data    = load_yaml("education.yaml")
    skills_data = load_yaml("skills.yaml")
    lang_data   = load_yaml("languages.yaml")
    cert_data   = load_yaml("certifications.yaml")
    proj_data   = load_yaml("projects.yaml")

    languages = [l["code"] for l in settings.get("languages_available", [{"code": "en"}])]

    for lang in languages:
        build_resume(lang, main_data, exp_data, edu_data,
                     skills_data, lang_data, cert_data, proj_data, settings)


if __name__ == "__main__":
    main()
